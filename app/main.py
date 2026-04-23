import asyncio
import json
import os
import re
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path
import queue as _tq

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from langgraph.types import Command

from app.agent import INITIAL_STATE, graph, _reg_stream, _unreg_stream

STATIC_DIR = Path(__file__).parent.parent / "static"

ASSETS_DIR = Path(__file__).parent.parent / "assets"

app = FastAPI(title="ResearchAgent")
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
app.mount("/assets", StaticFiles(directory=str(ASSETS_DIR)), name="assets")


@app.get("/")
async def root():
    return FileResponse(str(STATIC_DIR / "index.html"))


@app.get("/api/session")
async def create_session():
    return {"session_id": str(uuid.uuid4())}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_config(session_id: str) -> dict:
    return {"configurable": {"thread_id": session_id}}


def _get_interrupt_value(session_id: str) -> dict | None:
    state = graph.get_state(_get_config(session_id))
    if state.next:
        for task in state.tasks:
            for interrupt_obj in task.interrupts:
                return interrupt_obj.value
    return None


def _parse_markdown_table(markdown: str) -> list[list[str]]:
    rows = []
    for line in markdown.split("\n"):
        line = line.strip()
        if line.startswith("|") and line.endswith("|") and "---" not in line:
            cells = [c.strip() for c in line[1:-1].split("|")]
            rows.append(cells)
    return rows


STATUS_BY_TYPE = {
    "question": "Generando ecuación de búsqueda elaborada...",
    "equation": "Buscando artículos en ArXiv y Google Scholar...",
    "papers": "Generando matriz bibliográfica e hipótesis de investigación...",
    "qa": "Analizando tu pregunta...",
}


# ---------------------------------------------------------------------------
# XLSX download
# ---------------------------------------------------------------------------

HDR_FILL = PatternFill(start_color="DBEAFE", end_color="DBEAFE", fill_type="solid")
HDR_FONT = Font(color="1E40AF", bold=True)
ALT_FILL = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
ALT_FONT = Font(color="1E293B")
WRAP = Alignment(wrap_text=True, vertical="top")


def _set_header(ws, col: int, row: int, value: str):
    c = ws.cell(row=row, column=col, value=value)
    c.fill = HDR_FILL
    c.font = HDR_FONT
    c.alignment = WRAP


@app.get("/api/download/xlsx/{session_id}")
async def download_xlsx(session_id: str):
    config = _get_config(session_id)
    snap = graph.get_state(config)
    if not snap.values:
        return {"error": "Sesión no encontrada"}

    vals = snap.values
    papers = vals.get("papers", [])
    matrix_md = vals.get("matrix", "")
    equation = vals.get("search_equation", "")
    needs = vals.get("research_needs", "")
    paper_count = vals.get("paper_count", 5)

    wb = openpyxl.Workbook()

    # Sheet 1: Artículos
    ws1 = wb.active
    ws1.title = "Artículos"

    cols1 = [
        ("#", 4), ("Título", 48), ("Autores", 28), ("Año", 6),
        ("Fuente", 14), ("Revista / Conferencia", 30), ("DOI", 24),
        ("Citas", 8), ("Open Access", 12), ("Palabras Clave", 28),
        ("Volumen", 8), ("Número", 8), ("Páginas", 10), ("URL", 50),
    ]
    for i, (header, width) in enumerate(cols1, 1):
        _set_header(ws1, i, 1, header)
        ws1.column_dimensions[get_column_letter(i)].width = width

    for i, p in enumerate(papers, 1):
        row = i + 1
        values = [
            i, p.get("title", ""), p.get("authors", ""), p.get("year", ""),
            p.get("source", ""), p.get("journal", ""), p.get("doi", ""),
            p.get("citations", ""), p.get("open_access", ""), p.get("keywords", ""),
            p.get("volume", ""), p.get("issue", ""), p.get("pages", ""), p.get("url", ""),
        ]
        for col, val in enumerate(values, 1):
            cell = ws1.cell(row=row, column=col, value=val)
            cell.alignment = WRAP
            if i % 2 == 0:
                cell.fill = ALT_FILL
                cell.font = ALT_FONT

    # Sheet 2: Matriz
    ws2 = wb.create_sheet("Matriz Bibliográfica")
    if matrix_md:
        rows = _parse_markdown_table(matrix_md)
        for r, row_data in enumerate(rows, 1):
            for c, val in enumerate(row_data, 1):
                cell = ws2.cell(row=r, column=c, value=val)
                cell.alignment = WRAP
                if r == 1:
                    cell.fill = HDR_FILL
                    cell.font = HDR_FONT
                elif r % 2 == 0:
                    cell.fill = ALT_FILL
                    cell.font = ALT_FONT
        max_col = max((len(r) for r in rows), default=1)
        for c in range(1, max_col + 1):
            ws2.column_dimensions[get_column_letter(c)].width = 28

    # Sheet 3: Búsqueda
    ws3 = wb.create_sheet("Búsqueda")
    meta = [
        ("Tema de investigación", needs),
        ("Ecuación de búsqueda", equation),
        ("Papers por fuente", paper_count),
        ("Total artículos", len(papers)),
        ("ArXiv", sum(1 for p in papers if p.get("source") == "ArXiv")),
        ("Google Scholar", sum(1 for p in papers if p.get("source") == "Google Scholar")),
        ("Subidos", sum(1 for p in papers if p.get("source") == "Subido")),
        ("Fecha de búsqueda", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Modelo LLM", "llama3.1:8b (Ollama)"),
    ]
    for r, (k, v) in enumerate(meta, 1):
        ws3.cell(row=r, column=1, value=k).font = Font(bold=True)
        ws3.cell(row=r, column=2, value=str(v)).alignment = WRAP
    ws3.column_dimensions["A"].width = 28
    ws3.column_dimensions["B"].width = 65

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f"matriz_bibliografica_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ---------------------------------------------------------------------------
# DOCX download
# ---------------------------------------------------------------------------

@app.get("/api/download/docx/{session_id}")
async def download_docx(session_id: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return {"error": "python-docx no está instalado"}

    config = _get_config(session_id)
    snap = graph.get_state(config)
    if not snap.values:
        return {"error": "Sesión no encontrada"}

    vals = snap.values
    papers = vals.get("papers", [])
    matrix_md = vals.get("matrix", "")
    equation = vals.get("search_equation", "")
    explanation = vals.get("equation_explanation", "")
    needs = vals.get("research_needs", "")
    matrix_template = vals.get("matrix_template", "")

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def set_heading(para, text, level=1):
        para.style = f"Heading {level}"
        para.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        return para

    def add_heading(doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        return h

    # Cover
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Reporte de Investigación Bibliográfica")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Science Agent — DataHack 2026")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generado: {datetime.now().strftime('%d de %B de %Y, %H:%M')}")

    doc.add_page_break()

    # Section 1: Research topic
    add_heading(doc, "1. Tema de Investigación", 1)
    doc.add_paragraph(needs or "No especificado.")

    # Section 2: Methodology
    add_heading(doc, "2. Metodología", 1)
    doc.add_paragraph(
        "Este reporte fue generado utilizando Science Agent, un sistema de investigación "
        "bibliográfica basado en inteligencia artificial. El proceso siguió los siguientes pasos:"
    )

    steps = [
        ("Definición del tema", "El investigador describió su área de investigación o pregunta científica al agente conversacional."),
        ("Generación de ecuación de búsqueda", "Un modelo de lenguaje large (LLM) local mediante Ollama (llama3.1:8b) construyó una ecuación booleana elaborada con sinónimos, variantes morfológicas y operadores de campo específicos."),
        ("Búsqueda paralela", "El sistema consultó simultáneamente ArXiv (repositorio de preprints científicos) y Google Scholar, recuperando artículos con sus metadatos completos: título, autores, año, abstract, DOI, citas, revista y acceso abierto."),
        ("Selección de plantilla", "El investigador eligió una plantilla de matriz bibliográfica adaptada a su tipo de revisión (Estado del Arte, Revisión Sistemática, Benchmarking, Marco Teórico, Tendencias o Meta-análisis)."),
        ("Generación de matriz", "El LLM procesó todos los artículos encontrados y generó una matriz bibliográfica estructurada en Markdown, incluyendo análisis de tendencias dominantes, brechas identificadas y recomendaciones de lectura."),
        ("Análisis y exportación", "Los resultados fueron exportados a este documento Word con la descripción metodológica completa y los metadatos de todos los artículos."),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # Section 3: Search equation
    add_heading(doc, "3. Ecuación de Búsqueda", 1)
    doc.add_paragraph("La siguiente ecuación booleana fue utilizada para la recuperación de artículos:")

    eq_para = doc.add_paragraph()
    eq_run = eq_para.add_run(equation or "No disponible")
    eq_run.font.name = "Courier New"
    eq_run.font.size = Pt(9)
    eq_para.paragraph_format.left_indent = Cm(1)

    if explanation:
        add_heading(doc, "Explicación de la ecuación", 2)
        doc.add_paragraph(explanation)

    # Section 4: Papers
    add_heading(doc, "4. Artículos Recuperados", 1)

    arxiv_n = sum(1 for p in papers if p.get("source") == "ArXiv")
    scholar_n = sum(1 for p in papers if p.get("source") == "Google Scholar")
    uploaded_n = sum(1 for p in papers if p.get("source") == "Subido")

    doc.add_paragraph(
        f"Se recuperaron {len(papers)} artículos en total: {arxiv_n} de ArXiv, "
        f"{scholar_n} de Google Scholar y {uploaded_n} subidos manualmente."
    )

    for i, p in enumerate(papers, 1):
        add_heading(doc, f"{i}. {p.get('title', 'Sin título')}", 2)

        meta_lines = []
        if p.get("authors"):
            meta_lines.append(("Autores", p["authors"]))
        if p.get("year"):
            meta_lines.append(("Año", str(p["year"])))
        if p.get("source"):
            meta_lines.append(("Fuente", p["source"]))
        if p.get("journal"):
            meta_lines.append(("Revista", p["journal"]))
        if p.get("doi"):
            meta_lines.append(("DOI", p["doi"]))
        if p.get("citations"):
            meta_lines.append(("Citas", str(p["citations"])))
        if p.get("open_access"):
            meta_lines.append(("Acceso Abierto", p["open_access"]))
        if p.get("url"):
            meta_lines.append(("URL", p["url"]))

        for label, value in meta_lines:
            p_meta = doc.add_paragraph()
            run_l = p_meta.add_run(f"{label}: ")
            run_l.bold = True
            p_meta.add_run(value)

        if p.get("abstract"):
            doc.add_paragraph(f"Resumen: {p['abstract'][:600]}{'...' if len(p.get('abstract','')) > 600 else ''}")

        if p.get("keywords"):
            kw_p = doc.add_paragraph()
            kw_r = kw_p.add_run("Palabras clave: ")
            kw_r.bold = True
            kw_p.add_run(p["keywords"])

    # Section 5: Matrix
    if matrix_md:
        doc.add_page_break()
        add_heading(doc, "5. Matriz Bibliográfica", 1)

        template_names = {
            "estado_arte": "Estado del Arte",
            "sistematica": "Revisión Sistemática",
            "benchmarking": "Benchmarking Técnico",
            "marco_teorico": "Marco Teórico",
            "tendencias": "Tendencias e Innovación",
            "metaanalisis": "Meta-análisis",
        }
        tpl_name = template_names.get(matrix_template, matrix_template or "Personalizada")
        doc.add_paragraph(f"Plantilla utilizada: {tpl_name}")

        rows = _parse_markdown_table(matrix_md)
        if rows:
            try:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_val in enumerate(row_data):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_val
                        if r_idx == 0:
                            for par in cell.paragraphs:
                                for run in par.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "1D4ED8")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:val"), "clear")
                            tcPr.append(shd)
            except Exception:
                doc.add_paragraph("(La tabla no pudo ser generada en este formato)")
                doc.add_paragraph(matrix_md[:3000])

        # Extract trend analysis section
        trend_match = re.search(r'##\s*Análisis General(.*?)(?=^##|\Z)', matrix_md, re.DOTALL | re.MULTILINE)
        if trend_match:
            add_heading(doc, "6. Análisis de Tendencias", 1)
            trend_text = trend_match.group(1).strip()
            for line in trend_text.split("\n"):
                line = line.strip()
                if line.startswith("###"):
                    add_heading(doc, line.lstrip("#").strip(), 2)
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    p.add_run(line.strip("*")).bold = True
                elif line.startswith("-") or line.startswith("*"):
                    doc.add_paragraph(line.lstrip("-* "), style="List Bullet")
                elif line:
                    doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"metodologia_investigacion_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ---------------------------------------------------------------------------
# Paper upload
# ---------------------------------------------------------------------------

async def _extract_pdf_text(contents: bytes) -> str:
    try:
        import pdfplumber
        buf = BytesIO(contents)
        text_parts = []
        with pdfplumber.open(buf) as pdf:
            for page in pdf.pages[:8]:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"Error extracting PDF: {e}"


async def _extract_metadata_llm(text: str) -> dict:
    try:
        from langchain_ollama import ChatOllama
        llm = ChatOllama(
            model=os.getenv("OLLAMA_MODEL", "llama3.1:8b"),
            temperature=0,
        )
        prompt = (
            "Extract metadata from this academic paper. Return ONLY valid JSON with these exact keys:\n"
            '{"title":"","authors":"Author A, Author B","year":"2024","abstract":"","keywords":"kw1,kw2","journal":"","doi":""}\n\n'
            f"Paper text:\n{text[:2500]}\n\nJSON only:"
        )
        response = llm.invoke(prompt)
        raw = response.content.strip()
        start = raw.find("{")
        end = raw.rfind("}") + 1
        if start >= 0 and end > start:
            raw = raw[start:end]
        data = json.loads(raw)
        data["source"] = "Subido"
        data["url"] = ""
        data["pdf_url"] = ""
        data["open_access"] = "Sí"
        data["citations"] = ""
        return data
    except Exception:
        snippet = text[:400].replace("\n", " ")
        return {
            "title": "Paper subido manualmente",
            "authors": "",
            "year": str(datetime.now().year),
            "abstract": snippet,
            "keywords": "",
            "journal": "",
            "doi": "",
            "source": "Subido",
            "url": "",
            "pdf_url": "",
            "open_access": "Sí",
            "citations": "",
        }


@app.post("/api/upload/paper/{session_id}")
async def upload_paper(
    session_id: str,
    file: UploadFile | None = File(None),
    metadata: str | None = Form(None),
):
    config = _get_config(session_id)
    snap = graph.get_state(config)

    paper_data: dict = {}

    if metadata:
        try:
            paper_data = json.loads(metadata)
        except Exception:
            paper_data = {}

    if file and file.filename:
        contents = await file.read()
        if file.filename.lower().endswith(".pdf"):
            text = await _extract_pdf_text(contents)
            paper_data = await _extract_metadata_llm(text)
        else:
            return {"error": "Solo se admiten archivos PDF"}

    if not paper_data.get("title"):
        return {"error": "No se pudo extraer metadata del paper"}

    # Ensure required fields
    paper_data.setdefault("source", "Subido")
    paper_data.setdefault("url", "")
    paper_data.setdefault("pdf_url", "")
    paper_data.setdefault("citations", "")
    paper_data.setdefault("open_access", "Sí")

    # Update session state
    if snap.values:
        papers = list(snap.values.get("papers", []))
        papers.append(paper_data)
        graph.update_state(config, {"papers": papers})
    else:
        return {"error": "Sesión no encontrada. Inicia una búsqueda primero."}

    return {"success": True, "paper": paper_data}


# ---------------------------------------------------------------------------
# WebSocket
# ---------------------------------------------------------------------------

@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: str):
    await websocket.accept()
    config = _get_config(session_id)

    try:
        existing = graph.get_state(config)
        if existing.values:
            iv = _get_interrupt_value(session_id)
            if iv:
                await websocket.send_json(iv)
        else:
            await asyncio.to_thread(graph.invoke, dict(INITIAL_STATE), config)
            iv = _get_interrupt_value(session_id)
            if iv:
                await websocket.send_json(iv)

        while True:
            data = await websocket.receive_json()

            current_iv = _get_interrupt_value(session_id)
            current_type = (current_iv or {}).get("type", "")

            status_msg = STATUS_BY_TYPE.get(current_type)
            if status_msg:
                await websocket.send_json({"type": "status", "content": status_msg})

            if current_type == "equation":
                resume_value = {
                    "equation": data.get("content", "confirmar"),
                    "count": data.get("count", 5),
                }
            elif current_type == "papers":
                resume_value = {
                    "template_id": data.get("template_id", ""),
                    "custom": data.get("content", ""),
                }
            else:
                resume_value = data.get("content", "").strip()
                if not resume_value:
                    continue

            # ── Matrix streaming ────────────────────────────────────────────
            if current_type == "papers":
                loop = asyncio.get_running_loop()
                stream_q: asyncio.Queue = asyncio.Queue()
                _reg_stream(session_id, stream_q, loop)

                invoke_task = asyncio.create_task(
                    asyncio.to_thread(graph.invoke, Command(resume=resume_value), config)
                )

                while True:
                    try:
                        chunk = await asyncio.wait_for(stream_q.get(), timeout=3.0)
                        if chunk is None:
                            break
                        await websocket.send_json({"type": "matrix_chunk", "content": chunk})
                    except asyncio.TimeoutError:
                        if invoke_task.done():
                            break

                _unreg_stream(session_id)
                await websocket.send_json({"type": "status", "content": "Generando hipótesis de investigación..."})
                await invoke_task
            else:
                await asyncio.to_thread(graph.invoke, Command(resume=resume_value), config)
            # ────────────────────────────────────────────────────────────────

            new_iv = _get_interrupt_value(session_id)
            if new_iv:
                await websocket.send_json(new_iv)
            else:
                await websocket.send_json({"type": "complete", "content": "Sesión completada."})

    except WebSocketDisconnect:
        pass
    except Exception as exc:
        try:
            await websocket.send_json({"type": "error", "content": str(exc)})
        except Exception:
            pass
